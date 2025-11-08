from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls
import subprocess
import os
from positions import work_experience_positions, skills_data, education_data
try:
    from pypdf import PdfWriter, PdfReader
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False


class Resume:
    def __init__(self):
        self.doc = Document()
        self.setup_styles()
    
    def setup_styles(self):
        """Set up document styles and margins"""
        self.doc.styles['Normal'].font.name = "Times New Roman"
        self.doc.styles['Normal'].font.size = Pt(11)
        self.set_margins()
    
    def set_margins(self):
        """Set document margins - reduced left and right margins"""
        section = self.doc.sections[0]
        section.left_margin = Inches(0.5)    # Reduced from default 1 inch
        section.right_margin = Inches(0.5)   # Reduced from default 1 inch
        section.top_margin = Inches(1.0)     # Keep default
        section.bottom_margin = Inches(1.0)  # Keep default
    
    def add_header(self, name, contact_info):
        """Add header with name and contact information"""
        header = self.doc.sections[0].header
        
        # Main name
        p = header.add_paragraph(name)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(30)

        p.paragraph_format.space_after = Pt(0)
        
        # Contact information
        contact_p = header.add_paragraph(contact_info)
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in contact_p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
    
    def add_body_item(self, title, content=None, bullets=None):
        """Add a body item with title and optional content/bullets"""
        # Add title
        if title:
            heading = self.doc.add_paragraph()
            heading_run = heading.add_run(title.upper())
            heading_run.bold = True
            heading_run.font.size = Pt(14)
        
        # Add content
        if content:
            self.doc.add_paragraph(content)
        
        # Add bullets
        if bullets:
            for bullet in bullets:
                para = self.doc.add_paragraph(bullet)
                para.style = self.doc.styles['List Bullet']
    
    def add_work_experience(self, company, role, dates, bullets):
        """Add work experience section"""
        # Company name
        p = self.doc.add_paragraph()
        company_run = p.add_run(company)
        company_run.bold = True
        company_run.font.size = Pt(13)
        
        # Reduce space after company name paragraph
        p.paragraph_format.space_after = Pt(0)
        
        # Role and dates on separate line
        role_para = self.doc.add_paragraph()
        role_run = role_para.add_run(role)
        role_run.italic = True
        role_run.bold = True
        
        # Create a tab stop at the right margin for dates
        tab_stops = role_para.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT)
        
        # Use tab to push dates to the right
        role_para.add_run("\t")
        dates_run = role_para.add_run(dates)
        
        # Reduce space after role/dates paragraph
        role_para.paragraph_format.space_after = Pt(0)
        
        # Bullet points
        for bullet in bullets:
            para = self.doc.add_paragraph(bullet)
            para.style = self.doc.styles['List Bullet']
    
    def add_bottom_line(self, paragraph):
        """Add a bottom line to a paragraph (box with only bottom border visible)"""
        p_element = paragraph._element
        pPr = p_element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        
        # Add invisible top border
        top = OxmlElement('w:top')
        top.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'none')
        top.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz', '0')
        top.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}space', '0')
        top.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color', 'auto')
        pBdr.append(top)
        
        # Add invisible left border
        left = OxmlElement('w:left')
        left.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'none')
        left.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz', '0')
        left.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}space', '0')
        left.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color', 'auto')
        pBdr.append(left)
        
        # Add invisible right border
        right = OxmlElement('w:right')
        right.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'none')
        right.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz', '0')
        right.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}space', '0')
        right.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color', 'auto')
        pBdr.append(right)
        
        # Add visible bottom border
        bottom = OxmlElement('w:bottom')
        bottom.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'single')
        bottom.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz', '6')
        bottom.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}space', '1')
        bottom.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color', 'auto')
        pBdr.append(bottom)
        
        pPr.append(pBdr)
    
    def add_summary(self, summary_text):
        """Add summary section"""
        self.doc.add_paragraph(summary_text)
    
    def add_work_experience_title(self):
        """Add work experience section title with bottom line"""
        p = self.doc.add_paragraph()
        p_run = p.add_run("RELEVANT WORK EXPERIENCE")
        p_run.bold = True
        p_run.font.size = Pt(14)
        self.add_bottom_line(p)
    
    def add_work_experience_items(self):
        """Add all work experience items from positions data"""
        for position in work_experience_positions:
            self.add_work_experience(
                position["company"],
                position["title"], 
                position["dates"],
                position["items"]
            )
            
    def add_skills_title(self):
        """Add skills section title with bottom line"""
        # Add "SKILLS" title with bottom line
        p = self.doc.add_paragraph()
        p_run = p.add_run("SKILLS")
        p_run.bold = True
        p_run.font.size = Pt(14)
        self.add_bottom_line(p)
    
    def add_skills_items(self):
        """Add all skills items from skills data"""
        for i, skill_section in enumerate(skills_data):
            skill_para = self.doc.add_paragraph()
            skill_title = skill_para.add_run(f"{skill_section['section_title']}: ")
            skill_title.bold = True
            skill_para.add_run(skill_section['skills'])
            
            # Only add space_after = Pt(0) for all sections except the last one
            if i < len(skills_data) - 1:
                skill_para.paragraph_format.space_after = Pt(0)
    
    def add_education_title(self):
        """Add education section title with bottom line"""
        # Add "EDUCATION" title with bottom line
        p = self.doc.add_paragraph()
        p_run = p.add_run("EDUCATION")
        p_run.bold = True
        p_run.font.size = Pt(14)
        self.add_bottom_line(p)
    
    def add_education_items(self):
        """Add all education items from education data"""
        for i, education_item in enumerate(education_data):
            self.add_work_experience(
                education_item["institution"],
                education_item["degree"],
                education_item["dates"],
                ""
            )
            
            # Add spacing between items (except for the last one)
            if i < len(education_data) - 1:
                spacing_para = self.doc.add_paragraph()
                spacing_para.paragraph_format.space_after = Pt(1)
                for run in spacing_para.runs:
                    run.font.size = Pt(6)
    
    def add_body(self):
        """Add all body sections (skills, education, etc.)"""
        # Skills section with new format
        self.add_skills_title()
        self.add_skills_items()
        
        # Education section with new format
        self.add_education_title()
        self.add_education_items()
    
    def save(self, filename):
        """Save the resume to file"""
        self.doc.save(filename)
        print(f"Resume saved to: {filename}")
    
    def save_as_pdf(self, docx_filename, pdf_filename=None):
        """Convert DOCX to PDF using python-docx2pdf (pure Python solution)"""
        if pdf_filename is None:
            pdf_filename = docx_filename.replace('.docx', '.pdf')
        
        if not PYPDF_AVAILABLE:
            print("pypdf not available. Please install pypdf: pip install pypdf")
            return False
        
        try:
            # Use python-docx2pdf for DOCX to PDF conversion
            from docx2pdf import convert
            
            # Convert DOCX to PDF
            convert(docx_filename, pdf_filename)
            
            # Use pypdf to add metadata or process the PDF
            with open(pdf_filename, 'rb') as file:
                pdf_reader = PdfReader(file)
                pdf_writer = PdfWriter()
                
                # Copy all pages
                for page in pdf_reader.pages:
                    pdf_writer.add_page(page)
                
                # Add metadata
                pdf_writer.add_metadata({
                    "/Title": "Shimon Pozin - Resume",
                    "/Author": "Shimon Pozin",
                    "/Creator": "Resume Builder",
                    "/Subject": "Professional Resume"
                })
                
                # Write the final PDF
                with open(pdf_filename, 'wb') as output_file:
                    pdf_writer.write(output_file)
            
            print(f"PDF saved to: {pdf_filename}")
            return True
                
        except ImportError:
            print("docx2pdf not found. Installing...")
            try:
                import subprocess
                subprocess.check_call(['pip', 'install', 'docx2pdf'])
                from docx2pdf import convert
                convert(docx_filename, pdf_filename)
                print(f"PDF saved to: {pdf_filename}")
                return True
            except Exception as e:
                print(f"Failed to install docx2pdf: {e}")
                return False
        except Exception as e:
            print(f"PDF conversion error: {e}")
            return False
    
    def save_both_formats(self, base_filename):
        """Save resume in both DOCX and PDF formats"""
        docx_filename = f"{base_filename}.docx"
        pdf_filename = f"{base_filename}.pdf"
        
        # Save DOCX
        self.save(docx_filename)
        
        # Convert to PDF
        if self.save_as_pdf(docx_filename, pdf_filename):
            print(f"Both DOCX and PDF formats saved successfully!")
            return True
        else:
            print(f"DOCX saved, but PDF conversion failed.")
            if not PYPDF_AVAILABLE:
                print("To enable PDF conversion, install pypdf: pip install pypdf")
            print(f"You can manually convert {docx_filename} to PDF using LibreOffice or other tools.")
            return False


# Only run this code when the file is executed directly, not when imported
if __name__ == '__main__':
    # Create resume instance
    resume = Resume()

    # Add header
    resume.add_header("Shimon Pozin", "shimon.pozin@gmail.com | (416) 276-7917 | www.linkedin.com/in/shimonpozin")


    # Add summary
    resume.add_summary("I am a Business & Systems Analyst with 8+ years of experience delivering enterprise solutions across finance, insurance, B2B, and multimedia domains. I hold a Master of Computer Science and a PMP certification (since 2010). I am skilled in the full BA/SA lifecycle, including requirements elicitation, workflow modeling (BPMN/UML), BRD/SRS creation, API definition (REST/SOAP), and SDLC support. Furthermore, I am an expert in user stories, use cases, and acceptance criteria, and collaborating with cross-functional teams to align systems with business goals.")

    # Add work experience section
    resume.add_work_experience_title()
    resume.add_work_experience_items()

    # Add body sections
    resume.add_body()

    # Save in both DOCX and PDF formats
    output_base = "Shimon_Pozin_generated_resume"
    resume.save_both_formats(output_base)
